"""
word_exporter.py

Exporta DataFrames para Word (DOCX) usando python-docx.

Autor: Matheus
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_word(df: pd.DataFrame, output_path: str, title: str = "Relatório de Dados"):
    """
    Exporta um DataFrame para Word (DOCX) com formatação profissional.
    
    Args:
        df: DataFrame a ser exportado
        output_path: Caminho do arquivo DOCX de saída
        title: Título do relatório
    """
    
    # Limita a 100 linhas para evitar documentos gigantes
    MAX_ROWS = 100
    if len(df) > MAX_ROWS:
        df = df.head(MAX_ROWS)
        truncated = True
    else:
        truncated = False
    
    # Cria documento
    doc = Document()
    
    # Título
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do DataFrame
    info_text = f"Total de linhas: {len(df)} | Total de colunas: {len(df.columns)}"
    if truncated:
        info_text += f" (Exibindo apenas as primeiras {MAX_ROWS} linhas)"
    
    info_paragraph = doc.add_paragraph(info_text)
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espaço
    doc.add_paragraph()
    
    # Cria tabela
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        header_cells[i].text = str(col_name)
        
        # Formata header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Cor de fundo do header (azul escuro)
        shading_elm = header_cells[i]._element.get_or_add_tcPr()
        shading = shading_elm.get_or_add_shd()
        shading.fill = "2C3E50"
    
    # Dados
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text = str(value)[:100]  # Limita tamanho
            row_cells[i].text = cell_text
            
            # Formata células
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    # Salva documento
    doc.save(output_path)
    
    print(f"✅ Word exportado com sucesso: {output_path}")